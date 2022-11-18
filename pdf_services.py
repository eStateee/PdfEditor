from typing import AnyStr
from pdf2docx import parse
import docx
from docx2pdf import convert
from docx.shared import Pt
from constants import *
from PyPDF2 import PdfMerger, PdfWriter, PdfReader


def convert_pdf_to_docx() -> AnyStr:
    _doc_file_name = 'd.docx'
    pdf_path = os.path.join(TEMPLATES_PATH, 'before.pdf')
    docx_path = os.path.join(TEMP_PATH, _doc_file_name)
    parse(pdf_path, docx_path)
    return docx_path


def change_before_pdf_file(values, descr_type) -> AnyStr:
    name = values['-PROD_NAME-']
    global_prod_name = values['-GLOBAL_NAME1-']
    _D = {
        'qwe': name,
        'asd': global_prod_name.upper(),
        'zxc': descr_type.upper(),
    }
    docx_path = convert_pdf_to_docx()  # path to docx file
    docx_document = docx.Document(docx_path)  # docx Entity
    try:
        _redacted_pdf_file_path = redact_docx_file(docx_document, _D)
        return _redacted_pdf_file_path
    except Exception:
        print('Oops... Some exception:(')


def add_contacts(file_path, descr_type, values, filename) -> None:
    merger = PdfMerger()
    original_pdf_file = open(os.path.join(file_path), 'rb')
    redacted_pdf_file_path = change_before_pdf_file(values=values, descr_type=descr_type)
    before = open(os.path.join(redacted_pdf_file_path), 'rb')
    after = open(os.path.join(TEMPLATES_PATH, 'after.pdf'), 'rb')
    merger.append(original_pdf_file)
    merger.merge(position=0, fileobj=before, pages=(0, 1))
    merger.append(after)
    output = open(os.path.join(READY_PATH, filename), "wb")
    merger.write(output)
    merger.close()
    output.close()


def add_meta(filename, author, meta) -> PdfWriter:
    reader = PdfReader(os.path.join(READY_PATH, filename))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Author": author,
            "/Producer": "",
            '/Subject': meta,
            '/Title': meta
        }
    )
    return writer


def redact_docx_file(doc, D):
    _pdf_path = os.path.join(TEMP_PATH, 'doc.pdf')

    # REDACT
    for i in D.keys():
        for j in doc.paragraphs:
            if j.text.find(i) >= 0:
                j.text = j.text.replace(i, '')
                runner = j.add_run(D[i])
                font = runner.font
                font.name = 'Times New Roman'
                runner.font.size = docx.shared.Pt(36)
                runner.bold = True
                if i == 'qwe':
                    runner.underline = True
                elif i == 'asd':
                    runner.font.size = docx.shared.Pt(14)
                elif i == 'zxc':
                    runner.font.size = docx.shared.Pt(24)
                    runner.italic = True
    try:
        _redacted_docx_file_path = os.path.join(TEMP_PATH, 'doc.docx')
        doc.save(_redacted_docx_file_path)
        convert(_redacted_docx_file_path)  # convert docx to pdf(doc.pdf)
        return _pdf_path
    except Exception:
        return None


def write_new_pdf_file(filename, writer):
    with open(os.path.join(READY_PATH, filename), "wb") as file:
        writer.write(file)
    print(f'Файл {filename} Успешно изменен')
